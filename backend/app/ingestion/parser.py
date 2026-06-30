"""
Document parser for the ingestion pipeline.

PDF extraction uses PyMuPDF (fitz) with table-aware extraction:
- Tables are detected and serialised row-by-row into readable "Label — Header: value"
  text, so a policy table like the leave matrix stays semantically intact for retrieval
  (pypdf used to mash all columns into one unreadable line).
- Non-table prose is extracted separately, skipping repeated page headers/footers.
- A cleanup pass fixes mis-encoded characters (the PDF font maps apostrophes / hyphens to
  the U+FFFD replacement char) and de-hyphenates words broken across lines.

DOCX extraction also serialises tables. Scanned/image PDFs (no extractable text) are
detected and logged — OCR remains a separate concern (see app/ingestion/ocr.py).
"""
import os
import re
from typing import List, Dict
from app.core.logging import logger

DOCX_WORDS_PER_PAGE = 400

# Fraction of page height treated as header / footer margin (page numbers, running titles).
_MARGIN_RATIO = 0.06
_SHORT_MARGIN_CHARS = 70


def parse_document(file_path: str, content_type: str) -> List[Dict]:
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf" or content_type == "application/pdf":
        return _parse_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return [{"page": 1, "text": _clean_text(f.read())}]
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# --------------------------------------------------------------------------- #
# Text cleanup
# --------------------------------------------------------------------------- #
_REPLACEMENTS = {
    "‘": "'", "’": "'",       # curly single quotes
    "“": '"', "”": '"',       # curly double quotes
    "–": "-", "—": "-",       # en / em dash
    "•": "-",                       # bullet
    "\xa0": " ",                         # non-breaking space
}


def _clean_text(text: str) -> str:
    if not text:
        return ""
    # Soft hyphen (\xad) and the replacement char (�, a glyph the PDF font failed to
    # map) are almost always a hyphen/apostrophe dropped where a word wrapped across a line
    # or table cell. When one sits between word fragments, join them (consuming the spaces
    # the wrap left behind): "calculat\xad ed" -> "calculated", "al\xad lowed" -> "allowed".
    text = re.sub(r"(?<=\w)[ \t]*[\xad�][ \t]*(?=\w)", "", text)
    for bad, good in _REPLACEMENTS.items():
        text = text.replace(bad, good)
    text = text.replace("\xad", "")       # any leftover soft hyphens
    text = text.replace("�", " ")    # any leftover replacement chars
    # De-hyphenate words split across line breaks: "carry-\nforward" -> "carryforward"
    text = re.sub(r"-\s*\n\s*", "", text)
    # Normalise whitespace while preserving paragraph/line structure.
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _clean_cell(cell) -> str:
    return _clean_text((cell or "").replace("\n", " ")).strip()


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #
def _parse_pdf(path: str) -> List[Dict]:
    try:
        import pymupdf  # PyMuPDF
    except ImportError:  # pragma: no cover - defensive fallback
        return _parse_pdf_pypdf(path)

    pages: List[Dict] = []
    empty_pages = 0
    doc = pymupdf.open(path)
    try:
        for i, page in enumerate(doc):
            page_num = i + 1
            parts: List[str] = []
            table_rects = []

            try:
                tables = page.find_tables()
                for t in tables.tables:
                    serialised = _serialize_table(t)
                    if serialised:
                        parts.append(serialised)
                        table_rects.append(t.bbox)
            except Exception as e:
                logger.warning("table_extraction_failed", page=page_num, error=str(e))

            prose = _extract_prose(page, table_rects)
            if prose:
                parts.append(prose)

            page_text = _clean_text("\n\n".join(p for p in parts if p))
            if page_text:
                pages.append({"page": page_num, "text": page_text})
            else:
                empty_pages += 1
    finally:
        doc.close()

    # Heuristic: a PDF that yields almost no text is likely scanned -> needs OCR.
    if pages and empty_pages > len(pages):
        logger.warning("pdf_mostly_empty_possibly_scanned", path=os.path.basename(path),
                       empty_pages=empty_pages, text_pages=len(pages))
    if not pages:
        logger.warning("pdf_no_text_extracted", path=os.path.basename(path))

    return pages


def _extract_prose(page, table_rects) -> str:
    """Return page text outside table regions, skipping header/footer margins."""
    try:
        page_height = page.rect.height
        top_limit = page_height * _MARGIN_RATIO
        bottom_limit = page_height * (1 - _MARGIN_RATIO)
        blocks = page.get_text("blocks")  # (x0, y0, x1, y1, text, block_no, block_type)
        lines = []
        for b in blocks:
            x0, y0, x1, y1, text = b[0], b[1], b[2], b[3], b[4]
            if not text or not text.strip():
                continue
            if _inside_any(b[:4], table_rects):
                continue
            # Skip short blocks sitting in the top/bottom margin (page numbers, titles).
            stripped = text.strip()
            if (y1 < top_limit or y0 > bottom_limit) and len(stripped) < _SHORT_MARGIN_CHARS:
                continue
            lines.append(stripped)
        return "\n".join(lines)
    except Exception as e:
        logger.warning("prose_extraction_failed", error=str(e))
        return page.get_text() or ""


def _inside_any(block_bbox, table_rects) -> bool:
    bx0, by0, bx1, by1 = block_bbox
    bcx, bcy = (bx0 + bx1) / 2, (by0 + by1) / 2
    for r in table_rects:
        rx0, ry0, rx1, ry1 = r
        if rx0 <= bcx <= rx1 and ry0 <= bcy <= ry1:
            return True
    return False


def _serialize_table(table) -> str:
    """Turn a detected table into readable rows: 'RowLabel — Header: value; ...'."""
    try:
        rows = table.extract()
    except Exception:
        return ""
    if not rows:
        return ""

    header = [_clean_cell(c) for c in rows[0]]
    header_is_label = any(header)
    data_rows = rows[1:] if header_is_label else rows

    lines = []
    for row in data_rows:
        cells = [_clean_cell(c) for c in row]
        if not any(cells):
            continue
        label = cells[0]
        pairs = []
        for h, v in zip(header[1:], cells[1:]):
            if not v:
                continue
            pairs.append(f"{h}: {v}" if h else v)

        if label and pairs:
            lines.append(f"{label} — " + "; ".join(pairs))
        elif label:
            lines.append(label)
        elif pairs:
            lines.append("; ".join(pairs))
    return "\n".join(lines)


def _parse_pdf_pypdf(path: str) -> List[Dict]:
    """Last-resort fallback if PyMuPDF is unavailable."""
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages = []
    for i, page in enumerate(reader.pages):
        text = _clean_text(page.extract_text() or "")
        if text:
            pages.append({"page": i + 1, "text": text})
    return pages


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
def _parse_docx(path: str) -> List[Dict]:
    from docx import Document
    doc = Document(path)

    blocks: List[str] = []
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            blocks.append(para.text.strip())
    for table in doc.tables:
        serialised = _serialize_docx_table(table)
        if serialised:
            blocks.append(serialised)

    full_text = _clean_text("\n".join(blocks))
    words = full_text.split()
    if not words:
        return []

    pages = []
    page_num = 1
    for start in range(0, len(words), DOCX_WORDS_PER_PAGE):
        chunk_text = " ".join(words[start:start + DOCX_WORDS_PER_PAGE])
        if chunk_text.strip():
            pages.append({"page": page_num, "text": chunk_text})
            page_num += 1
    return pages


def _serialize_docx_table(table) -> str:
    rows = [[_clean_cell(c.text) for c in row.cells] for row in table.rows]
    if not rows:
        return ""
    header = rows[0]
    lines = []
    for row in rows[1:]:
        if not any(row):
            continue
        label = row[0]
        pairs = [f"{h}: {v}" if h else v for h, v in zip(header[1:], row[1:]) if v]
        if label and pairs:
            lines.append(f"{label} — " + "; ".join(pairs))
        elif label:
            lines.append(label)
        elif pairs:
            lines.append("; ".join(pairs))
    return "\n".join(lines)
